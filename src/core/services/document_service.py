"""Document processing service."""
import os
import logging
from pathlib import Path
from typing import Optional, Dict, List
from datetime import datetime

import PyPDF2
from werkzeug.datastructures import FileStorage

from mevzuat import db
from mevzuat.core.models import Document, DocumentPage
from mevzuat.core.utils.file_utils import allowed_file, get_file_extension

logger = logging.getLogger(__name__)

class DocumentService:
    """Service for document processing and management."""
    
    def __init__(self):
        self.supported_extensions = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'doc': self._process_doc,
            'txt': self._process_txt,
        }
    
    def save_document(self, file: FileStorage, metadata: Optional[Dict] = None) -> Document:
        """Save an uploaded document and create a database record.
        
        Args:
            file: The uploaded file
            metadata: Additional document metadata
            
        Returns:
            Document: The created document record
        """
        if not file or not allowed_file(file.filename):
            raise ValueError('Invalid file type')
        
        # Generate a unique filename
        file_ext = get_file_extension(file.filename)
        unique_id = str(uuid.uuid4())
        filename = f"{unique_id}.{file_ext}"
        
        # Save the file
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Create document record
        document = Document(
            title=metadata.get('title', file.filename),
            original_filename=file.filename,
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            file_type=file_ext,
            mime_type=file.mimetype,
            document_type=metadata.get('document_type'),
            document_number=metadata.get('document_number'),
            is_processed=False
        )
        
        db.session.add(document)
        db.session.commit()
        
        return document
    
    def process_document(self, document_id: int) -> bool:
        """Process a document by ID.
        
        Args:
            document_id: ID of the document to process
            
        Returns:
            bool: True if processing was successful, False otherwise
        """
        document = Document.query.get_or_404(document_id)
        
        try:
            # Check if already processed
            if document.is_processed:
                logger.info(f'Document {document_id} is already processed')
                return True
                
            # Get the appropriate processor for the file type
            file_ext = get_file_extension(document.original_filename).lower()
            processor = self.supported_extensions.get(file_ext)
            
            if not processor:
                raise ValueError(f'Unsupported file type: {file_ext}')
            
            # Process the document
            pages = processor(document.file_path)
            
            # Save document pages
            for page_num, (text, metadata) in enumerate(pages, 1):
                doc_page = DocumentPage(
                    document_id=document.id,
                    page_number=page_num,
                    text=text,
                    ocr_confidence=metadata.get('confidence'),
                    is_scanned=metadata.get('is_scanned', False)
                )
                db.session.add(doc_page)
            
            # Update document status
            document.is_processed = True
            document.processed_at = datetime.utcnow()
            document.processing_error = None
            
            db.session.commit()
            return True
            
        except Exception as e:
            logger.error(f'Error processing document {document_id}: {str(e)}', exc_info=True)
            document.processing_error = str(e)
            db.session.commit()
            return False
    
    def _process_pdf(self, file_path: str) -> List[tuple]:
        """Process a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of (text, metadata) tuples for each page
        """
        pages = []
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text() or ''
                    
                    # Simple heuristic to detect scanned documents
                    is_scanned = len(text.strip()) < 50  # Not much text found
                    
                    pages.append((text, {
                        'is_scanned': is_scanned,
                        'confidence': 1.0  # Not applicable for direct text extraction
                    }))
                    
        except Exception as e:
            logger.error(f'Error processing PDF {file_path}: {str(e)}')
            raise
            
        return pages
    
    def _process_docx(self, file_path: str) -> List[tuple]:
        """Process a DOCX file."""
        try:
            import docx
            doc = docx.Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs if para.text])
            return [(text, {'is_scanned': False})]
        except Exception as e:
            logger.error(f'Error processing DOCX {file_path}: {str(e)}')
            raise
    
    def _process_doc(self, file_path: str) -> List[tuple]:
        """Process a DOC file."""
        # Convert to text using antiword or similar tool
        try:
            import subprocess
            result = subprocess.run(['antiword', file_path], 
                                  capture_output=True, 
                                  text=True,
                                  check=True)
            return [(result.stdout, {'is_scanned': False})]
        except Exception as e:
            logger.error(f'Error processing DOC {file_path}: {str(e)}')
            raise
    
    def _process_txt(self, file_path: str) -> List[tuple]:
        """Process a text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return [(text, {'is_scanned': False})]
        except Exception as e:
            logger.error(f'Error processing TXT {file_path}: {str(e)}')
            raise
