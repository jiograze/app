"""
Belge işleme ana sınıfı - PDF, Word, vs. dosyaları işler
"""

import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pdfplumber

# PDF işleme
import PyPDF2

try:
    import fitz  # PyMuPDF

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Word işleme
from docx import Document as DocxDocument

# OCR (opsiyonel)
try:
    import pytesseract
    from PIL import Image

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from mevzuat.utils.document_classifier import DocumentClassifier
from mevzuat.utils.text_processor import TextProcessor
from .text_analyzer import TurkishTextAnalyzer


class DocumentProcessor:
    """Belge işleme ve analiz sınıfı"""

    def __init__(self, config_manager, database_manager):
        self.config = config_manager
        self.db = database_manager
        self.logger = logging.getLogger(self.__class__.__name__)

        # Alt bileşenler
        self.text_processor = TextProcessor(config_manager)
        self.document_classifier = DocumentClassifier(config_manager)
        self.text_analyzer = TurkishTextAnalyzer()

        # OCR ayarları
        self.ocr_enabled = config_manager.get("ocr.enabled", False)
        self.ocr_confidence_threshold = config_manager.get(
            "ocr.confidence_threshold", 75
        )
        self.ocr_language = config_manager.get("ocr.language", "tur+eng")

        # Önbellek
        self.document_cache = {}
        self.cache_enabled = config_manager.get("cache.enabled", True)
        self.cache_ttl = config_manager.get("cache.ttl_seconds", 3600)  # 1 saat

        # Desteklenen dosya türleri
        self.supported_extensions = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".doc": self._process_doc,
            ".odt": self._process_odt,
            ".txt": self._process_txt,
        }

        # OCR için gerekli kütüphanelerin kontrolü
        self._check_ocr_dependencies()
        
    def _check_ocr_dependencies(self):
        """OCR için gerekli kütüphanelerin kurulu olup olmadığını kontrol eder"""
        self.has_pillow = False
        self.has_pytesseract = False
        
        try:
            from PIL import Image
            self.has_pillow = True
            self.logger.debug("Pillow (PIL) kütüphanesi mevcut")
        except ImportError:
            self.logger.warning("Pillow (PIL) kütüphanesi kurulu değil. Görüntü işleme özellikleri devre dışı.")
            
        try:
            import pytesseract
            self.has_pytesseract = True
            self.logger.debug("Pytesseract kütüphanesi mevcut")
        except ImportError:
            self.logger.warning("Pytesseract kütüphanesi kurulu değil. OCR özellikleri devre dışı.")
            
        # OCR kullanılabilir mi?
        self.ocr_available = self.has_pillow and self.has_pytesseract
        if not self.ocr_available:
            self.logger.warning("OCR özellikleri için gerekli kütüphaneler eksik. PDF'lerde OCR işlemleri yapılamayacak.")
            
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Dosyanın hash'ini hesaplar
        
        Args:
            file_path: Hash'i hesaplanacak dosya yolu
            
        Returns:
            str: Dosyanın SHA-256 hash değeri
        """
        import hashlib
        
        # 64KB'lık bloklar halinde okuyup hash hesapla
        sha256_hash = hashlib.sha256()
        
        try:
            with open(file_path, "rb") as f:
                # Dosyayı parçalar halinde oku (büyük dosyalar için hafıza dostu)
                for byte_block in iter(lambda: f.read(65536), b""):
                    sha256_hash.update(byte_block)
            return sha256_hash.hexdigest()
        except Exception as e:
            self.logger.error(f"Hash hesaplama hatası ({file_path}): {str(e)}")
            # Hata durumunda dosya adı ve boyutunu kullanarak alternatif bir hash üret
            file_info = f"{file_path.name}_{file_path.stat().st_size}_{file_path.stat().st_mtime}"
            return hashlib.sha256(file_info.encode('utf-8')).hexdigest()

    def _get_cached_result(self, file_path: Path) -> Optional[Dict]:
        """Önbellekten sonuç getir"""
        if not self.cache_enabled:
            return None

        file_hash = self._calculate_file_hash(file_path)
        cache_entry = self.document_cache.get(file_hash)

        if cache_entry:
            # TTL kontrolü
            import time
            current_time = time.time()
            if current_time - cache_entry.get('timestamp', 0) < self.cache_ttl:
                self.logger.debug(f"Önbellekten yüklendi: {file_path}")
                return cache_entry['result']
            else:
                # Süresi dolmuş önbellek girdisini temizle
                del self.document_cache[file_hash]

        return None

    def _cache_result(self, file_path: Path, result: Dict):
        """Sonucu önbelleğe al"""
        if not self.cache_enabled:
            return

        file_hash = self._calculate_file_hash(file_path)
        self.document_cache[file_hash] = {
            'result': result,
            'timestamp': time.time(),
            'file_path': str(file_path)
        }
        self.logger.debug(f"Sonuç önbelleğe alındı: {file_path}")

    def process_file(self, file_path: str, use_cache: bool = True) -> Dict[str, Any]:
        """
        Ana dosya işleme fonksiyonu

        Args:
            file_path: İşlenecek dosya yolu
            use_cache: Önbellek kullanılsın mı?

        Returns:
            İşlem sonucu dict
        """
        try:
            file_path_obj = Path(file_path)
            self.logger.info(f"Dosya işleme başlatıldı: {file_path}")

            # Dosya varlık kontrolü
            if not file_path_obj.exists():
                return {"success": False, "error": "Dosya bulunamadı"}

            # Dosya boyutu kontrolü
            file_size = file_path_obj.stat().st_size
            max_size = self.config.get("max_file_size_mb", 100) * 1024 * 1024  # Varsayılan 100MB
            if file_size > max_size:
                return {
                    "success": False,
                    "error": f"Dosya boyutu çok büyük (Maksimum: {max_size / (1024 * 1024):.1f}MB)",
                }

            # Dosya uzantısı kontrolü
            ext = file_path_obj.suffix.lower()
            if ext not in self.supported_extensions:
                return {
                    "success": False,
                    "error": f"Desteklenmeyen dosya uzantısı: {ext}",
                }

            # Önbellek kontrolü
            if use_cache:
                cached_result = self._get_cached_result(file_path_obj)
                if cached_result:
                    return cached_result

            # Dosyayı işle
            process_func = self.supported_extensions[ext]
            result = process_func(file_path_obj)

            if not result["success"]:
                return result

            # Metin işleme
            raw_text = result["text"]
            if not raw_text.strip():
                return {"success": False, "error": "Dosyadan metin çıkarılamadı"}

            # Belge sınıflandırma
            classification_result = self.document_classifier.classify_document(
                raw_text, file_path_obj.name, result.get("metadata", {})
            )

            # Metin analizi ve maddeleme
            processed_data = self.text_processor.process_text(
                raw_text, classification_result
            )
            articles = processed_data.get("articles", [])

            # Veritabanına kaydet
            file_info = {
                "file_path": str(file_path_obj.absolute()),
                "file_size": file_size,
                "file_hash": self._calculate_file_hash(file_path_obj),
                "original_filename": file_path_obj.name,
                "file_type": ext[1:].upper(),  # Remove dot from extension
            }

            save_result = self._save_to_database(
                file_info, classification_result, articles, raw_text, result.get("metadata", {})
            )

            if save_result["success"]:
                # Dosya organizasyonu (eğer etkinse)
                organization_result = None
                if self.config.get("file_organization.enabled", True):
                    organization_result = self.organize_file(
                        str(file_path_obj), classification_result, file_info
                    )

                self.logger.info(f"Dosya başarıyla işlendi: {file_path}")
                result = {
                    "success": True,
                    "document_id": save_result["document_id"],
                    "title": classification_result.get("title", ""),
                    "document_type": classification_result.get("document_type", "DIGER"),
                    "articles_count": len(articles),
                    "classification": classification_result,
                    "file_info": file_info,
                    "text_length": len(raw_text),
                    "metadata": result.get("metadata", {}),
                    "processed_at": datetime.datetime.now().isoformat(),
                }

                # Organizasyon bilgilerini ekle
                if organization_result:
                    result["organization"] = organization_result

                # Sonucu önbelleğe al
                self._cache_result(file_path_obj, result)

                return result
            else:
                return save_result

        except Exception as e:
            self.logger.error(f"Dosya işleme hatası: {file_path} - {str(e)}", exc_info=True)
            return {"success": False, "error": str(e)}

    def _get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Dosya bilgilerini topla"""
        stat = file_path.stat()

        return {
            "original_filename": file_path.name,
            "file_path": str(file_path.absolute()),
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        }

    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """PDF dosyası işleme"""
        try:
            from PyPDF2 import PdfReader
            import tempfile
            import os

            text_content = ""
            has_scanned_pages = False
            total_pages = 0

            with open(file_path, "rb") as file:
                pdf = PdfReader(file)
                total_pages = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()

                    # Eğer sayfadan metin çıkarılamadıysa veya çok az metin varsa OCR dene
                    if not page_text or len(page_text.strip()) < 50:
                        if self.ocr_enabled and self.has_pillow and self.has_pytesseract:
                            has_scanned_pages = True

                            # Sayfayı görüntü olarak kaydet
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                                try:
                                    # PDF sayfasını görüntüye dönüştür
                                    from pdf2image import convert_from_path

                                    images = convert_from_path(
                                        str(file_path),
                                        first_page=page_num,
                                        last_page=page_num,
                                        dpi=300,
                                        fmt='png'
                                    )

                                    if images:
                                        images[0].save(temp_img.name, 'PNG')
                                        # OCR ile metni çıkar
                                        ocr_text = self._extract_text_with_ocr(temp_img.name)
                                        if ocr_text:
                                            text_content += f"\n[Sayfa {page_num} - TARANMIŞ BELGE]\n{ocr_text}\n\n"
                                except Exception as e:
                                    self.logger.warning(f"PDF sayfası OCR işleme hatası (sayfa {page_num}): {e}")
                                finally:
                                    # Geçici dosyayı temizle
                                    try:
                                        os.unlink(temp_img.name)
                                    except:
                                        pass
                    else:
                        text_content += f"\n[Sayfa {page_num}]\n{page_text}\n"

            metadata = {
                "pages": total_pages,
                "is_encrypted": pdf.is_encrypted,
                "has_scanned_content": has_scanned_pages,
                "info": pdf.metadata or {},
                "ocr_used": has_scanned_pages and self.ocr_enabled,
            }

            return {"success": True, "text": text_content, "metadata": metadata}

        except Exception as e:
            self.logger.error(f"PDF işleme hatası: {e}", exc_info=True)
            return {"success": False, "error": f"PDF işleme hatası: {e}"}

    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """DOCX dosyası işleme"""
        try:
            doc = DocxDocument(file_path)

            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            # Tabloları işle
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = ' '.join([p.text for p in cell.paragraphs])
                        row_text.append(cell_text)
                    text_content += "\t".join(row_text) + "\n"
                text_content += "\n"  # Tablolar arası boşluk

            # Resimlerden metin çıkarma (OCR ile)
            if self.ocr_enabled and self.has_pillow and self.has_pytesseract:
                import pytesseract
                from docx2python import docx2python
                from PIL import Image
                import io

                docx_content = docx2python(str(file_path))
                for rel in docx_content.images:
                    for img_path, img_data in docx_content.images[rel].items():
                        try:
                            # Resmi oku ve OCR uygula
                            image = Image.open(io.BytesIO(img_data))
                            ocr_text = pytesseract.image_to_string(image, lang=self.ocr_language)
                            if ocr_text.strip():
                                text_content += f"\n[RESİM METNİ]\n{ocr_text}\n"
                        except Exception as e:
                            self.logger.warning(f"Resim işleme hatası: {e}")

            metadata = {
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "images_count": len(doc.inline_shapes),
                "core_properties": {
                    "author": doc.core_properties.author or "",
                    "created": (
                        str(doc.core_properties.created)
                        if doc.core_properties.created
                        else ""
                    ),
                    "modified": (
                        str(doc.core_properties.modified)
                        if doc.core_properties.modified
                        else ""
                    ),
                    "last_modified_by": doc.core_properties.last_modified_by or "",
                    "revision": doc.core_properties.revision or "",
                    "version": doc.core_properties.version or "",
                },
                "has_ocr": self.ocr_enabled and self.has_pillow and self.has_pytesseract,
            }

            return {"success": True, "text": text_content, "metadata": metadata}

        except Exception as e:
            self.logger.error(f"DOCX işleme hatası: {e}", exc_info=True)
            return {"success": False, "error": f"DOCX işleme hatası: {e}"}

    def _process_doc(self, file_path: Path) -> Dict[str, Any]:
        """DOC dosyası işleme (eski format)"""
        # DOC formatı daha karmaşık, python-docx2txt veya antiword kullanılabilir
        # Şimdilik basit bir yaklaşım
        try:
            # Antiword varsa onu kullan
            import subprocess

            result = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True,
                text=True,
                encoding="utf-8",
            )

            if result.returncode == 0:
                return {
                    "success": True,
                    "text": result.stdout,
                    "metadata": {"extraction_method": "antiword"},
                }
            else:
                return {
                    "success": False,
                    "error": "DOC dosyası işlenemedi (antiword gerekli)",
                }

        except FileNotFoundError:
            return {
                "success": False,
                "error": "DOC dosyası işleme için antiword kurulu değil",
            }
        except Exception as e:
            return {"success": False, "error": f"DOC işleme hatası: {e}"}
            
    def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Düz metin dosyası işleme"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
                
            return {
                "success": True,
                "text": text_content,
                "metadata": {
                    "file_type": "text/plain",
                    "encoding": "utf-8"
                }
            }
            
        except UnicodeDecodeError:
            # UTF-8 başarısız olursa diğer kodlamaları dene
            try:
                with open(file_path, 'r', encoding='iso-8859-9') as f:  # Turkish encoding
                    text_content = f.read()
                    
                return {
                    "success": True,
                    "text": text_content,
                    "metadata": {
                        "file_type": "text/plain",
                        "encoding": "iso-8859-9"
                    }
                }
                
            except Exception as e:
                return {"success": False, "error": f"TXT dosyası okunamadı: {e}"}
                
        except Exception as e:
            return {"success": False, "error": f"TXT işleme hatası: {e}"}

    def _process_odt(self, file_path: Path) -> Dict[str, Any]:
        """ODT (OpenDocument Text) dosyası işleme"""
        try:
            from odf import text, teletype
            from odf.opendocument import load
            
            doc = load(str(file_path))
            text_content = ""
            
            # Ana metni çıkar
            for para in doc.getElementsByType(text.P):
                text_content += teletype.extractText(para) + "\n"
            
            # Tabloları işle
            for table in doc.getElementsByType(text.Table):
                for row in table.getElementsByType(text.TableRow):
                    row_text = []
                    for cell in row.getElementsByType(text.TableCell):
                        cell_text = teletype.extractText(cell)
                        row_text.append(cell_text.strip())
                    text_content += "\t".join(row_text) + "\n"
                text_content += "\n"  # Tablolar arası boşluk
            
            metadata = {
                "has_images": bool(doc.getElementsByType(text.Image)),
                "has_tables": bool(doc.getElementsByType(text.Table)),
            }
            
            # Meta verileri çıkar
            meta = doc.meta
            if meta:
                metadata.update({
                    "title": teletype.extractText(meta.getElementsByType(text.Title)[0]) if meta.getElementsByType(text.Title) else "",
                    "description": teletype.extractText(meta.getElementsByType(text.Description)[0]) if meta.getElementsByType(text.Description) else "",
                    "creator": teletype.extractText(meta.getElementsByType(text.Creator)[0]) if meta.getElementsByType(text.Creator) else "",
                    "date": teletype.extractText(meta.getElementsByType(text.Date)[0]) if meta.getElementsByType(text.Date) else "",
                })
            
            return {"success": True, "text": text_content, "metadata": metadata}
            
        except Exception as e:
            self.logger.error(f"ODT işleme hatası: {e}", exc_info=True)
            return {"success": False, "error": f"ODT işleme hatası: {e}"}

    def _compute_file_hash(self, file_path: Path) -> str:
        """Dosya MD5 hash'i hesapla"""
        try:
            import hashlib

            with open(file_path, "rb") as f:
                file_hash = hashlib.md5()
                chunk_size = 8192

                while chunk := f.read(chunk_size):
                    file_hash.update(chunk)

                return file_hash.hexdigest()

        except Exception as e:
            self.logger.error(f"Hash hesaplama hatası: {file_path} - {e}")
            return ""

    def _is_duplicate_file(self, file_hash: str) -> bool:
        """Hash'e göre duplicate dosya kontrolü"""
        if not file_hash:
            return False

        try:
            cursor = self.db.connection.cursor()
            cursor.execute(
                """
                SELECT id FROM documents WHERE file_hash = ? LIMIT 1
            """,
                (file_hash,),
            )

            result = cursor.fetchone()
            cursor.close()

            return result is not None

        except Exception as e:
            self.logger.error(f"Duplicate kontrol hatası: {e}")
            return False

    def _save_to_database(
        self,
        file_info: Dict,
        classification: Dict,
        articles: List[Dict],
        raw_text: str,
        metadata: Dict,
    ) -> Dict[str, Any]:
        """İşlenmiş belgeyi veritabanına kaydet"""
        try:
            # Dosya hash'i file_info'dan al (process_file'da zaten hesaplandı)
            file_hash = file_info.get("file_hash", "")

            # Belge kaydı
            document_data = {
                "title": classification.get("title", file_info["original_filename"]),
                "law_number": classification.get("law_number"),
                "document_type": classification.get("document_type", "DIGER"),
                "category": classification.get("category"),
                "subcategory": classification.get("subcategory"),
                "original_filename": file_info["original_filename"],
                "stored_filename": self._generate_stored_filename(
                    classification, file_info
                ),
                "file_path": file_info["file_path"],
                "file_hash": file_hash,
                "file_size": file_info["file_size"],
                "effective_date": classification.get("effective_date"),
                "publication_date": classification.get("publication_date"),
                "metadata": str(metadata),  # JSON string olarak
            }

            document_id = self.db.insert_document(document_data)

            # Madde kayıtları
            for idx, article in enumerate(articles):
                # Gelişmiş metin analizi ve temizleme
                text_analysis = self.text_analyzer.analyze_text(article["content"])
                fts_optimized_content = self.text_analyzer.prepare_for_fts(
                    article["content"]
                )

                article_data = {
                    "document_id": document_id,
                    "article_number": article.get("number"),
                    "title": article.get("title"),
                    "content": article["content"],
                    "content_clean": fts_optimized_content,  # Yeni gelişmiş temizlik
                    "seq_index": idx + 1,
                    "is_repealed": article.get("is_repealed", False),
                    "is_amended": article.get("is_amended", False),
                    "amendment_info": article.get("amendment_info"),
                    "article_type": article.get("type", "MADDE"),
                }

                self.db.insert_article(article_data)

            self.logger.info(
                f"Veritabanına kaydedildi: doc_id={document_id}, articles={len(articles)}"
            )

            return {"success": True, "document_id": document_id}

        except Exception as e:
            self.logger.error(f"Veritabanı kayıt hatası: {e}")
            return {"success": False, "error": f"Veritabanı kayıt hatası: {e}"}

    def _generate_stored_filename(self, classification: Dict, file_info: Dict) -> str:
        """Saklanacak dosya adını oluştur"""
        law_number = classification.get("law_number", "NA")
        title = classification.get("title", file_info["original_filename"])
        doc_type = classification.get("document_type", "DIGER")

        # Türkçe karakterleri normalize et ve slugify
        title_clean = self.text_processor.clean_text(title)
        title_slug = self.text_processor.slugify(title_clean)

        # Dosya uzantısı
        original_path = Path(file_info["original_filename"])
        extension = original_path.suffix

        # Dosya adı formatı: {numara}_{tip}_{baslik}.{uzanti}
        if law_number and law_number != "NA":
            # Numara varsa: 2022-4_genelge_dijitallesme.txt
            filename_base = f"{law_number}_{doc_type.lower()}_{title_slug[:30]}"
        else:
            # Numara yoksa: yonetmelik_cevre_koruma.txt
            filename_base = f"{doc_type.lower()}_{title_slug[:40]}"

        return f"{filename_base}{extension}"

    def _generate_organized_path(self, classification: Dict, file_info: Dict) -> Path:
        """Organize edilmiş dosya yolunu oluştur"""
        try:
            # Base depo klasörü
            base_folder = self.config.get_base_folder()
            organized_folder = base_folder / "organized"

            # Belge türüne göre ana klasör
            doc_type = classification.get("document_type", "DIGER")
            doc_type_folder = self._get_document_type_folder(doc_type)

            # Yıl ve numara bilgisi
            law_number = classification.get("law_number")
            pub_date = classification.get("publication_date")

            # Yıl çıkarma
            year = None
            if pub_date:
                try:
                    year = pub_date.split("-")[0]
                except:
                    pass

            # Dosya adından yıl çıkarma (genelge 2022/4 gibi)
            if not year:
                title = classification.get("title", "")
                filename = file_info["original_filename"]
                year_match = re.search(r"20\d{2}", f"{title} {filename}")
                if year_match:
                    year = year_match.group()

            # Yıl yoksa geçerli yıl
            if not year:
                year = str(datetime.now().year)

            # Klasör yapısını oluştur
            target_path = organized_folder / doc_type_folder

            # Yıl klasörü ekle
            if year:
                target_path = target_path / year

            # Numara klasörü ekle (genelge/yönetmelik için)
            if law_number or doc_type in ["YÖNERGE", "YÖNETMELİK"]:
                # Genelge numarası çıkarma (2022/4 formatı)
                number = law_number
                if not number:
                    # Dosya adından numara çıkarma
                    title = classification.get("title", "")
                    filename = file_info["original_filename"]

                    # 2022/4, 2023-12, No:15 gibi formatları ara
                    number_patterns = [
                        r"(\d{4}[/-]\d+)",  # 2022/4, 2023-12
                        r"[Nn]o\s*:?\s*(\d+)",  # No:4, no 15
                        r"sayı\s*:?\s*(\d+)",  # Sayı: 25
                        r"(\d+)\s*[/-]\s*\d{4}",  # 4/2022
                    ]

                    for pattern in number_patterns:
                        match = re.search(pattern, f"{title} {filename}")
                        if match:
                            number = match.group(1)
                            break

                if number:
                    target_path = target_path / str(number)

            return target_path

        except Exception as e:
            self.logger.error(f"Organize yol oluşturma hatası: {e}")
            # Hata durumunda basit yapı
            base_folder = self.config.get_base_folder()
            return base_folder / "organized" / "DIGER"

    def _get_document_type_folder(self, doc_type: str) -> str:
        """Belge türüne göre klasör adını döndür"""
        folder_mapping = {
            "ANAYASA": "anayasa",
            "KANUN": "kanun",
            "KHK": "khk",
            "TÜZÜK": "tuzuk",
            "YÖNETMELİK": "yonetmelik",
            "YÖNERGE": "genelge",  # Genelge ve yönerge aynı klasörde
            "TEBLİĞ": "teblig",
            "KARAR": "karar",
            "DIGER": "diger",
        }

        return folder_mapping.get(doc_type, "diger")

    def organize_file(
        self, file_path: str, classification: Dict, file_info: Dict
    ) -> Dict[str, Any]:
        """Dosyayı organize et ve uygun klasöre taşı"""
        try:
            source_path = Path(file_path)

            # Hedef yolu oluştur
            target_base = self._generate_organized_path(classification, file_info)

            # Klasörleri oluştur
            target_base.mkdir(parents=True, exist_ok=True)

            # Hedef dosya adı
            stored_filename = self._generate_stored_filename(classification, file_info)
            target_file = target_base / stored_filename

            # Eğer aynı isimde dosya varsa numara ekle
            counter = 1
            original_target = target_file
            while target_file.exists():
                name_part = original_target.stem
                extension = original_target.suffix
                target_file = target_base / f"{name_part}_{counter:02d}{extension}"
                counter += 1

            # Dosyayı taşı (kopyala + sil)
            import shutil

            shutil.copy2(source_path, target_file)

            # Orijinal dosyayı sil (opsiyonel)
            if self.config.get("file_organization.delete_original", True):
                source_path.unlink()
                self.logger.info(f"Orijinal dosya silindi: {source_path}")

            self.logger.info(f"Dosya organize edildi: {source_path} -> {target_file}")

            return {
                "success": True,
                "target_path": str(target_file),
                "target_folder": str(target_base),
                "organized_structure": self._get_folder_structure(target_base),
            }

        except Exception as e:
            self.logger.error(f"Dosya organizasyon hatası: {file_path} - {e}")
            return {"success": False, "error": f"Dosya organizasyon hatası: {e}"}

    def _get_folder_structure(self, path: Path) -> str:
        """Klasör yapısını string olarak döndür"""
        try:
            base_folder = self.config.get_base_folder()
            relative_path = path.relative_to(base_folder)
            return str(relative_path).replace("\\", "/")
        except:
            return str(path)

    def get_processing_stats(self) -> Dict[str, Any]:
        """İşleme istatistiklerini döndür"""
        # TODO: İmplementasyon
        return {
            "total_processed": 0,
            "success_rate": 0.0,
            "average_processing_time": 0.0,
        }
